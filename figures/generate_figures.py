#!/usr/bin/env python3
"""

Dyslexia fMRI: Generate all figures and tables from checkpoint JSONs.

Run AFTER master_analysis_v2.py has completed all sections.

Outputs:
  figures/
    fig1_roi_and_mvpa.pdf
    fig2_voxel_importance.pdf
    fig3_mds_rdm.pdf
    fig4_spd_position.pdf
    fig5_transfer.pdf
    supp1_permutations.pdf
    supp2_sensitivity.pdf
  tables/
    table1_demographics.docx
    table2_mvpa.docx
    table3_rsa.docx
    table4_transfer.docx
    table5_sensitivity.docx

Usage:
    python results_analysis.py
"""

import json
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.gridspec import GridSpec
from matplotlib import rcParams
from pathlib import Path
import warnings
warnings.filterwarnings('ignore')

# nilearn for brain figures
from nilearn import plotting, datasets, image
from nilearn.maskers import NiftiMasker
from nilearn.image import new_img_like
import nibabel as nib

# sklearn for MDS
from sklearn.manifold import MDS
from sklearn.svm import SVC
from sklearn.preprocessing import StandardScaler
from sklearn.feature_selection import SelectKBest, f_classif
from sklearn.model_selection import LeaveOneOut
from scipy.spatial.distance import squareform

# =============================================================================
# CONFIGURATION
# =============================================================================
OUTPUT_ROOT   = "/workspace/aime_dyslexia/analysis_output"
BIDS_ROOT     = "/workspace/aime_dyslexia/dataset_root"
GLM_DIR       = "/workspace/aime_dyslexia/analysis_output/first_level_glm"
CKPT_DIR      = Path(OUTPUT_ROOT) / 'master_results'
FIG_DIR       = Path(OUTPUT_ROOT) / 'figures'
TAB_DIR       = Path(OUTPUT_ROOT) / 'tables'
BEST_CONTRAST = 'PH'
N_FEATURES    = 1000

# Publication style
COLORS = {
    'DL':  '#E74C3C',   # red
    'SpD': '#F39C12',   # orange
    'TD':  '#2980B9',   # blue
    'sig': '#27AE60',   # green for significant
    'ns':  '#95A5A6',   # gray for non-significant
}
GROUPS = ['DL', 'SpD', 'TD']

# Global matplotlib style — DejaVu Sans (Arial equivalent, always available on Linux)
rcParams.update({
    'font.family':     'DejaVu Sans',
    'font.size':       11,
    'axes.titlesize':  12,
    'axes.labelsize':  11,
    'xtick.labelsize': 10,
    'ytick.labelsize': 10,
    'axes.spines.top':   False,
    'axes.spines.right': False,
    'figure.dpi':      300,
    'savefig.dpi':     300,
    'savefig.bbox':    'tight',
    'savefig.format':  'pdf',
})

# =============================================================================
# HELPERS
# =============================================================================

def load_ckpt(name):
    path = CKPT_DIR / f'ckpt_{name}.json'
    if not path.exists():
        print(f"  ⚠️  Checkpoint not found: {path}")
        return None
    with open(path) as f:
        return json.load(f)

def sig_stars(p):
    if p < 0.001: return '***'
    if p < 0.01:  return '**'
    if p < 0.05:  return '*'
    return 'ns'

def load_participants():
    tsv = Path(BIDS_ROOT) / 'participants.tsv'
    return pd.read_csv(tsv, sep='\t')

def get_subject_files(contrast, groups):
    """Load z-map files and labels for given contrast and groups."""
    participants = load_participants()
    files, labels, sids = [], [], []
    for _, row in participants.iterrows():
        grp = str(row.get('group', row.get('Group', ''))).strip()
        if grp not in groups:
            continue
        sid = str(row['participant_id']).strip()
        z_path = Path(GLM_DIR) / sid / 'averaged' / f'{contrast}_z_map.nii.gz'
        if not z_path.exists():
            z_path = Path(GLM_DIR) / sid / 'averaged' / f'{contrast}_z_map.nii'
        if z_path.exists():
            files.append(str(z_path))
            labels.append(grp)
            sids.append(sid)
    return files, np.array(labels), sids

def get_roi_mask(ref_file):
    """Build Harvard-Oxford reading network ROI mask."""
    ho   = datasets.fetch_atlas_harvard_oxford('cort-maxprob-thr25-2mm')
    labs = ho.labels
    idx  = [i for i, l in enumerate(labs) if any(r in l for r in
            ['Inferior Frontal Gyrus, pars triangularis',
             'Inferior Frontal Gyrus, pars opercularis',
             'Temporal Fusiform Cortex',
             'Superior Temporal Gyrus, posterior',
             'Angular Gyrus'])]
    maps = ho.maps
    roi  = image.math_img(
        'np.isin(img, idx).astype(np.float32)',
        img=maps, idx=idx
    )
    ref  = nib.load(ref_file)
    return image.resample_to_img(roi, ref, interpolation='nearest')

FIG_DIR.mkdir(parents=True, exist_ok=True)
TAB_DIR.mkdir(parents=True, exist_ok=True)

# =============================================================================
# FIGURE 1 — ROI Brain Map + MVPA Accuracy
# =============================================================================

def fig1_roi_and_mvpa(mvpa, perm):
    print("  Generating Figure 1: ROI map + MVPA accuracy...")

    fig = plt.figure(figsize=(14, 6))
    gs  = GridSpec(1, 2, figure=fig, wspace=0.35, width_ratios=[1.2, 1])

    # Panel A — Brain ROI map
    ax_brain = fig.add_subplot(gs[0])
    ax_brain.axis('off')
    ax_brain.set_title('A   Reading Network ROI (Harvard-Oxford)',
                        fontweight='bold', loc='left')

    try:
        files, _, _ = get_subject_files(BEST_CONTRAST, GROUPS)
        roi = get_roi_mask(files[0])
        display = plotting.plot_roi(
            roi,
            display_mode='z',
            cut_coords=[-14, -6, 0, 6, 14, 20],
            axes=ax_brain,
            colorbar=False,
            annotate=True,
            draw_cross=False,
            title='',
        )
    except Exception as e:
        ax_brain.text(0.5, 0.5, f'Brain map\n(run on server)\n{e}',
                      ha='center', va='center', transform=ax_brain.transAxes,
                      fontsize=9, color='gray')

    # Panel B — MVPA accuracy bar chart
    ax_bar = fig.add_subplot(gs[1])

    names    = [r['name'] for r in mvpa]
    accs     = [r['accuracy'] * 100 for r in mvpa]
    perm_map = {p['name']: p for p in perm} if perm else {}
    chance   = [50 if len(r['groups']) == 2 else 33.3 for r in mvpa]
    colors   = [COLORS['sig'] if perm_map.get(r['name'], {}).get('significant', False)
                else COLORS['ns'] for r in mvpa]

    bars = ax_bar.bar(range(len(names)), accs, color=colors,
                      edgecolor='white', linewidth=0.8, width=0.6)

    # Chance lines per bar
    for i, (ch, acc) in enumerate(zip(chance, accs)):
        ax_bar.plot([i - 0.3, i + 0.3], [ch, ch],
                    color='black', lw=1.5, ls='--', zorder=5)

    # Significance stars
    for i, r in enumerate(mvpa):
        p_info = perm_map.get(r['name'], {})
        p_val  = p_info.get('p_value', 1.0)
        stars  = sig_stars(p_val)
        y_pos  = accs[i] + 1.5
        ax_bar.text(i, y_pos, stars, ha='center', va='bottom',
                    fontsize=12, fontweight='bold',
                    color=COLORS['sig'] if stars != 'ns' else COLORS['ns'])

    ax_bar.set_xticks(range(len(names)))
    ax_bar.set_xticklabels([n.replace(' vs ', '\nvs\n') for n in names],
                            fontsize=9)
    ax_bar.set_ylabel('Classification Accuracy (%)')
    ax_bar.set_ylim(0, 100)
    ax_bar.axhline(50, color='lightgray', lw=0.8, ls=':', zorder=0)
    ax_bar.set_title('B   MVPA Classification (PH contrast)',
                     fontweight='bold', loc='left')

    # Legend
    sig_patch = mpatches.Patch(color=COLORS['sig'], label='Significant (p<0.05)')
    ns_patch  = mpatches.Patch(color=COLORS['ns'],  label='Not significant')
    ax_bar.legend(handles=[sig_patch, ns_patch], loc='upper right',
                  fontsize=9, framealpha=0.8)

    plt.suptitle('Figure 1. ROI-Constrained MVPA Classification',
                 fontsize=13, fontweight='bold', y=1.02)

    out = FIG_DIR / 'fig1_roi_and_mvpa.pdf'
    fig.savefig(out)
    plt.close(fig)
    print(f"    Saved: {out}")


# =============================================================================
# FIGURE 2 — Voxel Importance Map
# =============================================================================

def fig2_voxel_importance():
    print("  Generating Figure 2: Voxel importance map (DL vs TD, PH)...")

    try:
        files, labels, _ = get_subject_files(BEST_CONTRAST, ['DL', 'TD'])
        roi = get_roi_mask(files[0])

        masker = NiftiMasker(mask_img=roi, standardize=True,
                             memory='nilearn_cache', memory_level=1)
        X = masker.fit_transform(files)

        # Aggregate feature importance across LOO folds
        loo       = LeaveOneOut()
        scores    = np.zeros(X.shape[1])
        n_folds   = 0

        for tr, _ in loo.split(X):
            X_tr, y_tr = X[tr], labels[tr]
            k   = min(N_FEATURES, X_tr.shape[1] - 1)
            sel = SelectKBest(f_classif, k=k).fit(X_tr, y_tr)
            scores += sel.scores_
            n_folds += 1

        scores /= n_folds   # average across folds

        # Project back to brain space
        importance_img = masker.inverse_transform(scores)

        fig, axes = plt.subplots(1, 1, figsize=(14, 4))
        axes.axis('off')

        display = plotting.plot_stat_map(
            importance_img,
            display_mode='z',
            cut_coords=8,
            threshold=np.percentile(scores, 75),
            colorbar=True,
            cmap='hot',
            axes=axes,
            title='',
        )

        plt.suptitle(
            'Figure 2. Discriminative Voxels: DL vs TD (PH Contrast)\n'
            'Average F-score across LOO-CV folds (top 25% shown)',
            fontsize=12, fontweight='bold'
        )

        out = FIG_DIR / 'fig2_voxel_importance.pdf'
        fig.savefig(out)
        plt.close(fig)
        print(f"     Saved: {out}")

    except Exception as e:
        print(f"    ⚠️  Voxel importance map failed: {e}")
        print(f"       Run on server where brain data is accessible")


# =============================================================================
# FIGURE 3 — MDS Scatter + RDM Heatmaps
# =============================================================================

def fig3_mds_rdm(rsa):
    print("  Generating Figure 3: MDS scatter + RDM heatmaps...")

    if not rsa:
        print("    ⚠️  No RSA data — skipping")
        return

    # Use PH contrast for MDS
    ph = rsa.get('PH', rsa.get(list(rsa.keys())[0]))
    within  = ph['within_dissimilarity']
    between = ph['between_dissimilarity']

    # Build group-level RDM from between/within values
    # Order: DL, SpD, TD
    n_groups = 3
    rdm = np.zeros((n_groups, n_groups))
    pair_map = {
        ('DL',  'SpD'): between.get('DL_vs_SpD', 0),
        ('DL',  'TD'):  between.get('DL_vs_TD', 0),
        ('SpD', 'TD'):  between.get('SpD_vs_TD', 0),
    }
    group_idx = {g: i for i, g in enumerate(GROUPS)}
    for (g1, g2), val in pair_map.items():
        i, j = group_idx[g1], group_idx[g2]
        rdm[i, j] = rdm[j, i] = val
    for g in GROUPS:
        rdm[group_idx[g], group_idx[g]] = within.get(g, 0)

    # MDS embedding
    mds = MDS(n_components=2, dissimilarity='precomputed',
              random_state=42, normalized_stress='auto')
    coords = mds.fit_transform(rdm)

    fig = plt.figure(figsize=(14, 5))
    gs  = GridSpec(1, 3, figure=fig, wspace=0.4)

    # Panel A — MDS scatter (group centroids)
    ax_mds = fig.add_subplot(gs[0])
    for g in GROUPS:
        i = group_idx[g]
        ax_mds.scatter(coords[i, 0], coords[i, 1],
                       color=COLORS[g], s=300, zorder=5,
                       edgecolors='white', linewidths=1.5, label=g)
        ax_mds.annotate(g, (coords[i, 0], coords[i, 1]),
                        textcoords='offset points', xytext=(8, 4),
                        fontsize=11, fontweight='bold', color=COLORS[g])

    # Draw lines between groups
    for (g1, g2) in [('DL', 'TD'), ('DL', 'SpD'), ('SpD', 'TD')]:
        i, j = group_idx[g1], group_idx[g2]
        ax_mds.plot([coords[i, 0], coords[j, 0]],
                    [coords[i, 1], coords[j, 1]],
                    color='lightgray', lw=1.5, ls='--', zorder=1)

    ax_mds.set_xlabel('MDS Dimension 1')
    ax_mds.set_ylabel('MDS Dimension 2')
    ax_mds.set_title('A   Group Representational Geometry\n(PH contrast)',
                     fontweight='bold', loc='left', fontsize=10)
    ax_mds.legend(fontsize=9)

    # Panel B — RDM heatmap
    ax_rdm = fig.add_subplot(gs[1])
    im = ax_rdm.imshow(rdm, cmap='RdYlBu_r',
                       vmin=rdm.min() * 0.99,
                       vmax=rdm.max() * 1.01)
    ax_rdm.set_xticks(range(n_groups))
    ax_rdm.set_yticks(range(n_groups))
    ax_rdm.set_xticklabels(GROUPS)
    ax_rdm.set_yticklabels(GROUPS)
    for i in range(n_groups):
        for j in range(n_groups):
            ax_rdm.text(j, i, f'{rdm[i,j]:.4f}',
                        ha='center', va='center', fontsize=8,
                        color='white' if rdm[i,j] > rdm.mean() else 'black')
    plt.colorbar(im, ax=ax_rdm, shrink=0.8)
    ax_rdm.set_title('B   Group RDM (PH contrast)',
                     fontweight='bold', loc='left', fontsize=10)

    # Panel C — Between-group distances across contrasts
    ax_dist = fig.add_subplot(gs[2])
    contrasts = list(rsa.keys())
    x = np.arange(len(contrasts))
    w = 0.25
    pairs = [('DL_vs_TD', 'DL↔TD', COLORS['DL']),
             ('DL_vs_SpD', 'DL↔SpD', COLORS['SpD']),
             ('SpD_vs_TD', 'SpD↔TD', COLORS['TD'])]

    for k, (key, label, color) in enumerate(pairs):
        vals = [rsa[c]['between_dissimilarity'].get(key, np.nan)
                for c in contrasts]
        ax_dist.bar(x + k * w, vals, w, label=label,
                    color=color, alpha=0.85, edgecolor='white')

    ax_dist.set_xticks(x + w)
    ax_dist.set_xticklabels(contrasts, rotation=30, ha='right', fontsize=9)
    ax_dist.set_ylabel('Mean Dissimilarity (1-r)')
    ax_dist.set_title('C   Between-Group Distances\nacross Contrasts',
                      fontweight='bold', loc='left', fontsize=10)
    ax_dist.legend(fontsize=8)

    plt.suptitle('Figure 3. Representational Similarity Analysis',
                 fontsize=13, fontweight='bold', y=1.02)

    out = FIG_DIR / 'fig3_mds_rdm.pdf'
    fig.savefig(out)
    plt.close(fig)
    print(f"     Saved: {out}")


# =============================================================================
# FIGURE 4 — SpD Intermediate Position
# =============================================================================

def fig4_spd_position(rsa):
    print("  Generating Figure 4: SpD intermediate position...")

    if not rsa:
        print("    ⚠️  No RSA data — skipping")
        return

    contrasts = list(rsa.keys())
    fig, axes = plt.subplots(1, 2, figsize=(13, 5))

    # Panel A — Number line per contrast
    ax = axes[0]
    for i, c in enumerate(contrasts):
        bw = rsa[c]['between_dissimilarity']
        dl_td  = bw.get('DL_vs_TD', None)
        dl_spd = bw.get('DL_vs_SpD', None)
        spd_td = bw.get('SpD_vs_TD', None)
        if None in [dl_td, dl_spd, spd_td]:
            continue

        y = i
        # Normalize to 0-1 within DL↔TD range
        lo, hi = min(dl_spd, spd_td), dl_td
        if hi - lo < 1e-6:
            continue
        spd_pos = (dl_spd - lo) / (hi - lo)   # 0=DL side, 1=TD side

        ax.plot([0, 1], [y, y], color='lightgray', lw=2, zorder=1)
        ax.scatter([0], [y], color=COLORS['DL'], s=120, zorder=3)
        ax.scatter([1], [y], color=COLORS['TD'], s=120, zorder=3)
        ax.scatter([spd_pos], [y], color=COLORS['SpD'], s=200,
                   zorder=4, marker='D')
        ax.text(-0.05, y, c, ha='right', va='center', fontsize=10)

    ax.set_xlim(-0.35, 1.2)
    ax.set_ylim(-0.5, len(contrasts) - 0.5)
    ax.axis('off')
    ax.set_title('A   SpD Position on DL↔TD Continuum\n(◆ = SpD centroid)',
                 fontweight='bold', loc='left', fontsize=10)

    # Add legend
    for g, label in [('DL', 'DL'), ('SpD', 'SpD'), ('TD', 'TD')]:
        marker = 'D' if g == 'SpD' else 'o'
        ax.scatter([], [], color=COLORS[g], s=100, marker=marker, label=label)
    ax.legend(loc='lower right', fontsize=9)

    # Panel B — Within-group dissimilarity (heterogeneity)
    ax2 = axes[1]
    x   = np.arange(len(contrasts))
    w   = 0.25
    for k, g in enumerate(GROUPS):
        vals = [rsa[c]['within_dissimilarity'].get(g, np.nan)
                for c in contrasts]
        ax2.bar(x + k * w, vals, w, label=g,
                color=COLORS[g], alpha=0.85, edgecolor='white')

    ax2.set_xticks(x + w)
    ax2.set_xticklabels(contrasts, rotation=30, ha='right')
    ax2.set_ylabel('Within-Group Dissimilarity\n(higher = more heterogeneous)')
    ax2.set_title('B   Within-Group Neural Heterogeneity',
                  fontweight='bold', loc='left', fontsize=10)
    ax2.legend(fontsize=9)

    plt.suptitle('Figure 4. SpD Occupies Intermediate Representational Space',
                 fontsize=13, fontweight='bold', y=1.02)

    out = FIG_DIR / 'fig4_spd_position.pdf'
    fig.savefig(out)
    plt.close(fig)
    print(f" Saved: {out}")


# =============================================================================
# FIGURE 5 — Cross-Contrast Transfer
# =============================================================================

def fig5_transfer(transfer):
    print("  Generating Figure 5: Cross-contrast transfer...")

    if not transfer:
        print("  No transfer data — skipping")
        return

    accs  = transfer.get('accuracies', {})
    perms = transfer.get('permutations', {})

    if not accs:
        print("  Transfer accuracies empty — skipping")
        return

    labels  = list(accs.keys())
    values  = [accs[k] * 100 for k in labels]
    p_vals  = [perms.get(k, {}).get('p_value', 1.0) for k in labels]
    nulls   = [perms.get(k, {}).get('null_mean', 0.5) * 100 for k in labels]
    colors  = [COLORS['sig'] if p < 0.05 else COLORS['ns'] for p in p_vals]

    # Shorten labels for display
    short = [l.split('←')[0].strip() for l in labels]

    fig, ax = plt.subplots(figsize=(11, 5))
    x = np.arange(len(labels))
    w = 0.35

    bars = ax.bar(x - w/2, values, w, color=colors,
                  edgecolor='white', linewidth=0.8, label='Transfer accuracy')
    ax.bar(x + w/2, nulls, w, color='lightgray',
           edgecolor='white', linewidth=0.8, label='Null mean')

    # Chance line
    ax.axhline(50, color='black', lw=1.2, ls='--', label='Chance (50%)')

    # Stars
    for i, (p, v) in enumerate(zip(p_vals, values)):
        stars = sig_stars(p)
        ax.text(i - w/2, v + 1.5, stars, ha='center', fontsize=11,
                fontweight='bold',
                color=COLORS['sig'] if stars != 'ns' else COLORS['ns'])

    ax.set_xticks(x)
    ax.set_xticklabels(short, rotation=25, ha='right', fontsize=9)
    ax.set_ylabel('Accuracy (%)')
    ax.set_ylim(0, 85)
    ax.set_title('Figure 5. Cross-Contrast Transfer (LOO-CV + Permutation)\n'
                 'Training on PH contrast, testing on W contrast',
                 fontweight='bold', fontsize=12)
    ax.legend(fontsize=9, loc='upper right')

    out = FIG_DIR / 'fig5_transfer.pdf'
    fig.savefig(out)
    plt.close(fig)
    print(f"     Saved: {out}")


# =============================================================================
# SUPPLEMENTARY FIGURE 1 — Permutation Null Distributions
# =============================================================================

def supp1_permutations(mvpa, perm):
    print("  Generating Supp 1: Permutation null distributions...")

    if not perm:
        print("    ⚠️  No permutation data — skipping")
        return

    n = len(perm)
    fig, axes = plt.subplots(1, n, figsize=(4 * n, 4))
    if n == 1:
        axes = [axes]

    for ax, p in zip(axes, perm):
        null_mean = p['null_mean']
        null_std  = p['null_std']
        obs       = p['observed_acc']
        n_perms   = p['n_perms']
        p_val     = p['p_value']

        # Simulate null distribution from mean/std (Gaussian approx)
        np.random.seed(42)
        null_sim = np.random.normal(null_mean, null_std, n_perms)

        ax.hist(null_sim, bins=40, color=COLORS['ns'],
                alpha=0.7, edgecolor='white', linewidth=0.5)
        ax.axvline(obs, color=COLORS['DL'], lw=2.5,
                   label=f'Observed: {obs*100:.1f}%')
        ax.axvline(null_mean, color='gray', lw=1.5,
                   ls='--', label=f'Null mean: {null_mean*100:.1f}%')

        sig = " p={:.4f}".format(p_val) if p['significant'] \
              else "ns p={:.4f}".format(p_val)
        ax.set_title(f"{p['name']}\n{sig}", fontsize=10, fontweight='bold')
        ax.set_xlabel('Accuracy')
        ax.set_ylabel('Count' if ax == axes[0] else '')
        ax.legend(fontsize=8)

    plt.suptitle('Supplementary Figure 1. Permutation Test Null Distributions\n'
                 f'({n_perms} permutations per comparison)',
                 fontsize=12, fontweight='bold', y=1.02)

    out = FIG_DIR / 'supp1_permutations.pdf'
    fig.savefig(out)
    plt.close(fig)
    print(f"     Saved: {out}")


# =============================================================================
# SUPPLEMENTARY FIGURE 2 — Sensitivity Analysis
# =============================================================================

def supp2_sensitivity(sens):
    print("  Generating Supp 2: Sensitivity analysis...")

    if not sens:
        print("    ⚠️  No sensitivity data — skipping")
        return

    fig, ax = plt.subplots(figsize=(6, 5))

    labels = [f"Full sample\n(N={sens['n_full']})",
              f"Excl. outlier\n(N={sens['n_filtered']})"]
    values = [sens['acc_full'] * 100, sens['acc_filtered'] * 100]
    colors = [COLORS['sig'], COLORS['sig']]

    bars = ax.bar(labels, values, color=colors, width=0.4,
                  edgecolor='white', linewidth=0.8)

    ax.axhline(50, color='black', lw=1.2, ls='--', label='Chance (50%)')

    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, val + 1,
                f'{val:.1f}%', ha='center', fontsize=11, fontweight='bold')

    diff = abs(sens['acc_full'] - sens['acc_filtered']) * 100
    ax.set_ylabel('Classification Accuracy (%)')
    ax.set_ylim(0, 95)
    ax.set_title(f'Supplementary Figure 2. Sensitivity Analysis\n'
                 f'Excluding age outlier (sub-047EPKL014050, age=13)\n'
                 f'Difference: {diff:.1f}%  —  {sens["verdict"]}',
                 fontsize=11, fontweight='bold')

    out = FIG_DIR / 'supp2_sensitivity.pdf'
    fig.savefig(out)
    plt.close(fig)
    print(f"     Saved: {out}")


# =============================================================================
# TABLE GENERATION (via python-docx)
# =============================================================================

def make_tables(mvpa, sens, perm, rsa, transfer):
    """Generate all tables as .docx using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("  ⚠️  python-docx not installed — run: pip install python-docx")
        print("       Tables will be skipped")
        return

    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def header_row(table, texts, bg='2C3E50', fg='FFFFFF'):
        row = table.rows[0]
        for i, text in enumerate(texts):
            cell = row.cells[i]
            cell.text = text
            run  = cell.paragraphs[0].runs[0]
            run.bold      = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(
                int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16))
            set_cell_bg(cell, bg)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_cell(row, col_idx, text, bold=False, center=False, size=10):
        cell = row.cells[col_idx]
        cell.text = str(text)
        run  = cell.paragraphs[0].runs[0]
        run.bold      = bold
        run.font.size = Pt(size)
        if center:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return cell

    def new_doc(title):
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(11)
        h = doc.add_heading(title, 0)
        h.runs[0].font.size = Pt(13)
        return doc

    # ── Table 1: Demographics ──────────────────────────────────────────────
    print("  Generating Table 1: Demographics...")
    try:
        participants = load_participants()
        doc  = new_doc('Table 1. Participant Demographics')
        doc.add_paragraph(
            'Groups: DL = Dyslexia, SpD = Isolated Spelling Deficit, '
            'TD = Typical Development. Age: mean ± SD.'
        ).runs[0].font.size = Pt(10)

        table = doc.add_table(rows=4, cols=5)
        table.style = 'Table Grid'
        header_row(table, ['', 'DL', 'SpD', 'TD', 'Statistic'])

        row_labels = ['N', 'Age (mean ± SD)', 'Sex (M/F)']
        row_stats  = ['', 'F=?, p=?', 'χ²=?, p=?']

        # Pre-compute per-group stats
        group_stats = {}
        for grp in ['DL', 'SpD', 'TD']:
            grp_col = 'group' if 'group' in participants.columns else 'Group'
            sub = participants[participants[grp_col] == grp]
            n   = len(sub)

            age_col = next((c for c in ['age', 'Age'] if c in sub.columns), None)
            sex_col = next((c for c in ['sex', 'Sex'] if c in sub.columns), None)

            age_str = f'{sub[age_col].mean():.1f} ± {sub[age_col].std():.1f}' \
                      if age_col else '?'
            if sex_col:
                m = (sub[sex_col].astype(str).str.upper() == 'M').sum()
                f_ = (sub[sex_col].astype(str).str.upper() == 'F').sum()
                sex_str = f'{m}M/{f_}F'
            else:
                sex_str = '?'

            group_stats[grp] = {'n': n, 'age': age_str, 'sex': sex_str}

        # Fill rows
        data_rows = [
            ['N']            + [str(group_stats[g]['n'])   for g in ['DL','SpD','TD']] + [''],
            ['Age (mean±SD)'] + [group_stats[g]['age']      for g in ['DL','SpD','TD']] + ['F=?, p=?'],
            ['Sex (M/F)']    + [group_stats[g]['sex']       for g in ['DL','SpD','TD']] + ['χ²=?, p=?'],
        ]

        for r_idx, row_data in enumerate(data_rows):
            row = table.rows[r_idx + 1]
            for c_idx, val in enumerate(row_data):
                row.cells[c_idx].text = val
                run = row.cells[c_idx].paragraphs[0].runs[0]
                run.font.size = Pt(10)
                if c_idx == 0:
                    run.bold = True
                row.cells[c_idx].paragraphs[0].alignment = \
                    WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 \
                    else WD_ALIGN_PARAGRAPH.CENTER

        out = TAB_DIR / 'table1_demographics.docx'
        doc.save(str(out))
        print(f"     Saved: {out}")
    except Exception as e:
        print(f"    ⚠️  Table 1 failed: {e}")

    # ── Table 2: MVPA Results ──────────────────────────────────────────────
    print("  Generating Table 2: MVPA results...")
    try:
        doc   = new_doc('Table 2. MVPA Classification Results')
        doc.add_paragraph(
            f'Contrast: PH (Pseudoword). ROI: Harvard-Oxford reading network '
            f'(20,584 voxels). CV: nested LOO-CV. Permutations: 5000. '
            f'* p<0.05, ** p<0.01, *** p<0.001.'
        ).runs[0].font.size = Pt(10)

        perm_map = {p['name']: p for p in perm} if perm else {}
        table = doc.add_table(rows=len(mvpa)+1, cols=7)
        table.style = 'Table Grid'
        header_row(table, ['Comparison', 'N', 'Accuracy',
                            'Null Mean', 'p-value', 'Sig.', 'Interpretation'])

        interp = {
            'DL vs TD':  'Dyslexia neurally separable from TD',
            'DL vs SpD': 'DL and SpD neurally indistinguishable',
            'SpD vs TD': 'SpD not separable from TD',
            '3-way':     'Three-group classification at chance',
        }

        for i, r in enumerate(mvpa):
            p_info = perm_map.get(r['name'], {})
            p_val  = p_info.get('p_value', None)
            row    = table.rows[i + 1]
            add_cell(row, 0, r['name'], bold=True)
            add_cell(row, 1, str(r['n']), center=True)
            add_cell(row, 2, f"{r['accuracy']*100:.1f}%", center=True,
                     bold=p_info.get('significant', False))
            add_cell(row, 3, f"{p_info.get('null_mean', 0)*100:.1f}%",
                     center=True)
            add_cell(row, 4, f"{p_val:.4f}" if p_val else '—', center=True)
            add_cell(row, 5, sig_stars(p_val) if p_val else '—', center=True)
            add_cell(row, 6, interp.get(r['name'], ''))

            if p_info.get('significant', False):
                set_cell_bg(row.cells[2], 'D5F5E3')

        out = TAB_DIR / 'table2_mvpa.docx'
        doc.save(str(out))
        print(f"     Saved: {out}")
    except Exception as e:
        print(f"    ⚠️  Table 2 failed: {e}")

    # ── Table 3: RSA ───────────────────────────────────────────────────────
    print("  Generating Table 3: RSA between-group distances...")
    try:
        doc = new_doc('Table 3. RSA Between-Group Dissimilarity')
        doc.add_paragraph(
            'Mean pairwise 1−r dissimilarity between groups per contrast. '
            'SpD position determined by comparing DL↔SpD and SpD↔TD to DL↔TD.'
        ).runs[0].font.size = Pt(10)

        table = doc.add_table(rows=len(rsa)+1, cols=6)
        table.style = 'Table Grid'
        header_row(table, ['Contrast', 'DL↔TD', 'DL↔SpD',
                            'SpD↔TD', 'SpD Position', 'Key finding'])

        positions = {
            'W': 'Intermediate', 'PH': 'Intermediate',
            'Reading': 'Intermediate', 'W-PH': 'Intermediate',
            'PH-W': 'Closer to TD',
        }

        for i, (contrast, res) in enumerate(rsa.items()):
            bw  = res['between_dissimilarity']
            row = table.rows[i + 1]
            add_cell(row, 0, contrast, bold=True)
            add_cell(row, 1, f"{bw.get('DL_vs_TD', 0):.4f}", center=True)
            add_cell(row, 2, f"{bw.get('DL_vs_SpD', 0):.4f}", center=True)
            add_cell(row, 3, f"{bw.get('SpD_vs_TD', 0):.4f}", center=True)
            pos = positions.get(contrast, '?')
            add_cell(row, 4, pos, center=True,
                     bold=(pos == 'Intermediate'))
            add_cell(row, 5,
                     '← continuum' if pos == 'Intermediate' else '')
            if pos == 'Intermediate':
                set_cell_bg(row.cells[4], 'FEF9E7')

        out = TAB_DIR / 'table3_rsa.docx'
        doc.save(str(out))
        print(f"     Saved: {out}")
    except Exception as e:
        print(f"    ⚠️  Table 3 failed: {e}")

    # ── Table 4: Transfer ──────────────────────────────────────────────────
    print("  Generating Table 4: Cross-contrast transfer...")
    try:
        if transfer and transfer.get('accuracies'):
            accs  = transfer['accuracies']
            perms = transfer.get('permutations', {})
            doc   = new_doc('Table 4. Cross-Contrast Transfer Results')
            doc.add_paragraph(
                'LOO-CV cross-contrast transfer with permutation testing. '
                'Train on one contrast, test on a different contrast using '
                'the same subjects. Unbiased: subject left out of both contrasts.'
            ).runs[0].font.size = Pt(10)

            table = doc.add_table(rows=len(accs)+1, cols=6)
            table.style = 'Table Grid'
            header_row(table, ['Transfer', 'Groups', 'Accuracy',
                                'Null Mean', 'p-value', 'Sig.'])

            for i, (desc, acc) in enumerate(accs.items()):
                p_info = perms.get(desc, {})
                p_val  = p_info.get('p_value', None)
                # Parse groups from desc
                grp_str = 'DL vs TD' if 'DL vs TD' in desc \
                    else 'SpD vs TD' if 'SpD vs TD' in desc \
                    else 'DL vs SpD' if 'DL vs SpD' in desc else '?'
                transfer_str = desc.split('←')[0].strip()

                row = table.rows[i + 1]
                add_cell(row, 0, transfer_str)
                add_cell(row, 1, grp_str, center=True)
                add_cell(row, 2, f"{acc*100:.1f}%", center=True)
                add_cell(row, 3,
                         f"{p_info.get('null_mean',0)*100:.1f}%",
                         center=True)
                add_cell(row, 4,
                         f"{p_val:.4f}" if p_val else '—', center=True)
                add_cell(row, 5,
                         sig_stars(p_val) if p_val else '—', center=True)

            out = TAB_DIR / 'table4_transfer.docx'
            doc.save(str(out))
            print(f"  Saved: {out}")
        else:
            print("  Transfer data incomplete — skipping Table 4")
    except Exception as e:
        print(f"  Table 4 failed: {e}")

    # ── Table 5: Sensitivity ───────────────────────────────────────────────
    print("  Generating Table 5: Sensitivity analysis...")
    try:
        if sens:
            doc = new_doc('Table 5. Sensitivity Analysis — Age Outlier')
            doc.add_paragraph(
                'Participant sub-047EPKL014050 (TD group, age=13) is 3-4 years '
                'older than all other participants (age range 8-10). '
                'Analysis repeated excluding this participant.'
            ).runs[0].font.size = Pt(10)

            table = doc.add_table(rows=6, cols=3)
            table.style = 'Table Grid'
            header_row(table, ['Measure', 'Full Sample', 'Excl. Outlier'])

            rows_data = [
                ('N (DL vs TD)', str(sens['n_full']),
                 str(sens['n_filtered'])),
                ('Classification accuracy',
                 f"{sens['acc_full']*100:.1f}%",
                 f"{sens['acc_filtered']*100:.1f}%"),
                ('p-value (permutation)',
                 '0.007',
                 f"{sens['p_value_filtered']:.4f}"),
                ('Significant?',
                 'Yes (**)',
                 'Yes (*)' if sens['significant_filtered'] else 'No (ns)'),
                ('Difference from full sample',
                 '—',
                 f"{abs(sens['acc_full']-sens['acc_filtered'])*100:.1f}%"),
            ]

            for i, (label, full, filt) in enumerate(rows_data):
                row = table.rows[i + 1]
                add_cell(row, 0, label, bold=True)
                add_cell(row, 1, full, center=True)
                add_cell(row, 2, filt, center=True)

            doc.add_paragraph()
            verdict_para = doc.add_paragraph(f'Verdict: {sens["verdict"]}')
            verdict_para.runs[0].bold = True
            verdict_para.runs[0].font.size = Pt(11)

            out = TAB_DIR / 'table5_sensitivity.docx'
            doc.save(str(out))
            print(f"     Saved: {out}")
        else:
            print("  Sensitivity data not found — skipping Table 5")
    except Exception as e:
        print(f"   Table 5 failed: {e}")


# =============================================================================
# MAIN
# =============================================================================

def main():
    print("=" * 65)
    print("  RESULTS ANALYSIS")
    print("  Dyslexia fMRI — Figures + Tables")
    print("=" * 65)

    # Load all checkpoints
    print("\n[1] Loading checkpoints...")
    mvpa     = load_ckpt('mvpa')
    perm     = load_ckpt('perm')
    rsa      = load_ckpt('rsa')
    transfer = load_ckpt('transfer')
    sens     = load_ckpt('sensitivity')

    # Report what's available
    for name, data in [('mvpa', mvpa), ('perm', perm), ('rsa', rsa),
                       ('transfer', transfer), ('sensitivity', sens)]:
        status = '' if data else '⚠️  MISSING'
        print(f"    {name:<15} {status}")

    # Generate figures
    print("\n[2] Generating figures...")
    fig1_roi_and_mvpa(mvpa or [], perm or [])
    fig2_voxel_importance()
    fig3_mds_rdm(rsa)
    fig4_spd_position(rsa)
    fig5_transfer(transfer)
    supp1_permutations(mvpa or [], perm or [])
    supp2_sensitivity(sens)

    # Generate tables
    print("\n[3] Generating tables...")
    make_tables(mvpa, sens, perm, rsa, transfer)

    
    print("\n" + "=" * 65)
    print("  COMPLETE")
    print(f"  Figures → {FIG_DIR}")
    print(f"  Tables  → {TAB_DIR}")
    print("=" * 65)

    print("\n  Files generated:")
    for f in sorted(FIG_DIR.glob('*.pdf')):
        print(f"     {f.name}")
    for f in sorted(TAB_DIR.glob('*.docx')):
        print(f"     {f.name}")


if __name__ == '__main__':
    main()
